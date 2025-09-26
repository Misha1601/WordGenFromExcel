import os
import sys
import subprocess
from pathlib import Path
from docx import Document
import openpyxl
from datetime import datetime
import configparser


def replace_in_paragraph(paragraph, old_text, new_text):
    """Заменяет текст в параграфе"""
    if old_text not in paragraph.text:
        return

    runs_text = [run.text for run in paragraph.runs]
    full_text = ''.join(runs_text)

    if old_text not in full_text:
        return

    new_full_text = full_text.replace(old_text, new_text)

    # Очищаем все runs
    for run in paragraph.runs:
        run.text = ''

    # Устанавливаем новый текст
    paragraph.runs[0].text = new_full_text


def replace_text_in_doc(doc, old_text, new_text):
    """Заменяет текст во всех элементах документа"""
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, old_text, new_text)
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for nested_paragraph in nested_cell.paragraphs:
                                replace_in_paragraph(nested_paragraph, old_text, new_text)

def get_document_text(doc):
    """Извлекает весь текст документа для проверки замен"""
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            text.append(nested_cell.text)
    return "\n".join(text)

def load_config():
    """Загружает конфигурацию из INI файла с валидацией значений."""

    config = configparser.ConfigParser()
    config_file_path = os.path.join(os.getcwd(), 'WordGenFromExcel.ini')

    if not os.path.exists(config_file_path):
        print("INI файл не найден!")
        input("Нажмите Enter для выхода ...")
        exit(1)
    try:
        config.read(config_file_path, encoding='utf-8')
        template_name = config.get('PATHS', 'template_name')
        data_file_name = config.get('PATHS', 'data_file_name')

        # валидация
        if not template_name or not isinstance(template_name, str):
            raise ValueError("Название файла шаблона договора в ini файле указано не корректно")
        if not data_file_name or not isinstance(data_file_name, str):
            raise ValueError("Название файла эксел с данными в ini файле указано не корректно")
        if not template_name.lower().endswith('.docx'):
            raise ValueError("Название файла с шаблоном договора должен оканчиваться на .docx")
        if not data_file_name.lower().endswith('.xlsx'):
            raise ValueError("Название файла с данными экселе должен оканчиваться на .xlsx")
        if not os.path.exists(os.path.join(os.getcwd(), template_name)):
            raise ValueError(f"Не найден файл шаблона договора - {template_name}")
        if not os.path.exists(os.path.join(os.getcwd(), data_file_name)):
            raise ValueError(f"Не найден файл эксел с данными - {data_file_name}")

        return template_name, data_file_name

    except Exception as e:
        print(f"Ошибка: {e}.")
        print("Проверьте ini файл на корректное заполнение или убедитесь в наличии файла шаблона и файла с данными!")
        input("Нажмите Enter для выхода ...")
        exit(1)


def main():
    # Загрузка конфигурации
    template_name, data_file_name = load_config()
    # Конфигурация путей
    exe_dir = os.getcwd()

    template_path = os.path.join(exe_dir, template_name)
    xlsx_path = os.path.join(exe_dir, data_file_name)

    try:
        # Работа с Excel-данными
        wb = openpyxl.load_workbook(xlsx_path)
        sheet = wb.active
        rows = list(sheet.iter_rows(values_only=True))

        if not rows:
            raise ValueError("Файл Excel не содержит данных")

        headers = []
        for cell in rows[0]:
            if cell is not None and str(cell).strip() != "":
                headers.append(str(cell))
            else:
                break
        columns_count = len(headers)
        if columns_count == 0:
            raise ValueError("Все ячейки верхней строки Excel файла должны быть заполнены!")

        for row_idx, row in enumerate(rows[1:], 1):
            # Нормализация данных
            row_data = []
            for cell in row[:columns_count]:
                if cell is None:
                    row_data.append("-")
                elif isinstance(cell, datetime):
                    row_data.append(cell.strftime("%d.%m.%Y"))
                else:
                    row_data.append(str(cell).strip() or "-")

            # Дополнение данных до количества столбцов
            row_data += ["-"] * (columns_count - len(row_data))

            # Извлечение имени документа
            doc_name = row_data[0].strip() or f"row_{row_idx}"
            if not doc_name:
                print(f"Предупреждение: Пустое имя в строке {row_idx}, пропуск")
                continue

            # Загрузка шаблона
            doc = Document(template_path)

            # Выполнение замен
            for col_idx in range(1, columns_count):
                placeholder = headers[col_idx]
                value = row_data[col_idx]
                replace_text_in_doc(doc, placeholder, value)
                print(f"Замена: {placeholder} → {value}")

            # Проверка незамененных плейсхолдеров
            doc_text = get_document_text(doc)
            missing = [ph for ph in headers[1:] if ph in doc_text]

            if missing:
                missing_str = ", ".join(missing)
                print(f"Незамененные значения в шаблоне: {missing_str}.")
                error_msg = (
                    "Незамененные значения в шаблоне: "
                    f"{missing_str}. Убедитесь в едином форматировании!"
                )

                ps_script = (
                    'Add-Type -AssemblyName PresentationFramework;'
                    f'[System.Windows.MessageBox]::Show("{error_msg}", "Ошибка")'
                )
                subprocess.run(
                    ["powershell", "-Command", ps_script],
                    check=False
                )

            # Сохранение результата
            output_file = f"{doc_name}{Path(template_name).suffix}"
            output_path = os.path.join(exe_dir, output_file)
            doc.save(output_path)
            print(f"Создан документ: {output_path}\n")

        wb.close()

    except Exception as e:
        print(f"КРИТИЧЕСКАЯ ОШИБКА: {str(e)}", file=sys.stderr)
        input("Нажмите Enter для выхода ...")
        sys.exit(1)

if __name__ == "__main__":
    main()