import os
import sys
import subprocess
from pathlib import Path
from docx import Document
import openpyxl
from datetime import datetime
import configparser


def replace_in_paragraph(paragraph, old_text, new_text):
    """
    Заменяет все вхождения old_text на new_text в одном параграфе.
    Поведение имитирует Microsoft Word.
    """
    # Защита от пустого old_text или пустого параграфа
    if not old_text or not paragraph.runs:
        return

    # Объединяем текст всех run'ов параграфа в одну строку для поиска
    full_text = ''.join(run.text for run in paragraph.runs)

    # Если искомого текста нет — выходим
    if old_text not in full_text:
        return

    # Начинаем поиск с позиции 0
    start_search = 0

    # будем искать все вхождения по очереди
    while True:
        # Ищем позицию первого вхождения old_text, начиная с start_search
        pos = full_text.find(old_text, start_search)

        # Если больше нет вхождений — выходим из цикла
        if pos == -1:
            break

        # Конец заменяемого фрагмента (не включительно)
        end_pos = pos + len(old_text)
        # Получаем актуальный список run'ов (он может меняться при удалении)
        runs = paragraph.runs
        # current_offset — текущая позиция в full_text, соответствующая началу текущего run
        current_offset = 0
        # Список для хранения информации о run'ах, затронутых заменой:
        # каждый элемент: (индекс_run, начало_сегмента_в_run, конец_сегмента_в_run)
        affected_runs = []

        # Проходим по всем run'ам, чтобы понять, какие из них пересекаются с [pos, end_pos)
        for i, run in enumerate(runs):
            run_len = len(run.text)
            run_start = current_offset          # позиция начала этого run в full_text
            run_end = current_offset + run_len  # позиция конца этого run в full_text

            # Проверяем, пересекается ли этот run с заменяемым фрагментом
            if run_end <= pos or run_start >= end_pos:
                pass # Run полностью до или после заменяемого фрагмента — пропускаем
            else:
                # Этот run частично или полностью входит в заменяемый фрагмент

                # Где внутри run начинается заменяемый фрагмент?
                seg_start = max(0, pos - run_start)
                # Где внутри run заканчивается заменяемый фрагмент?
                seg_end = min(run_len, end_pos - run_start)
                # Сохраняем информацию об этом run
                affected_runs.append((i, seg_start, seg_end))
            # Сдвигаем offset на длину текущего run
            current_offset = run_end

        # Если по какой-то причине не нашли затронутые run — пропускаем это вхождение
        if not affected_runs:
            start_search = pos + 1
            continue

        # Определяем первый и последний затронутые run
        first_idx = affected_runs[0][0]   # индекс первого run, где начинается old_text
        last_idx = affected_runs[-1][0]   # индекс последнего run, где заканчивается old_text

        # === Часть 1: текст ДО old_text в первом run ===
        # Берём всё, что в первом run идёт до начала old_text
        before = runs[first_idx].text[:affected_runs[0][1]]
         # === Часть 2: текст ПОСЛЕ old_text в последнем run ===
        # Берём всё, что в последнем run идёт после конца old_text
        after = runs[last_idx].text[affected_runs[-1][2]:]

        # === Главное действие: замена текста ===
        # В ПЕРВОМ затронутом run оставляем "до" + новый текст
        runs[first_idx].text = before + new_text

        # Обновляем последний run, если он не первый
        if first_idx != last_idx:
            runs[last_idx].text = after

        # Удаляем промежуточные run (между first и last)
        for idx in range(last_idx - 1, first_idx, -1):
            paragraph._element.remove(runs[idx]._element)

        # Обновляем для следующей итерации
        runs = paragraph.runs
        full_text = ''.join(run.text for run in runs)
        # Сдвигаем позицию поиска за пределы только что заменённого фрагмента
        start_search = pos + len(new_text)
        # Защита от бесконечного цикла, если old_text == new_text
        if old_text == new_text:
            start_search += 1


def replace_text_in_doc(doc, old_text, new_text):
    """Заменяет текст во всех элементах документа"""
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, old_text, new_text)
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for nested_paragraph in nested_cell.paragraphs:
                                replace_in_paragraph(nested_paragraph, old_text, new_text)

def get_document_text(doc):
    """Извлекает весь текст документа для проверки замен"""
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            text.append(nested_cell.text)
    return "\n".join(text)

def load_config():
    """Загружает конфигурацию из INI файла с валидацией значений."""

    config = configparser.ConfigParser()
    config_file_path = os.path.join(os.getcwd(), 'WordGenFromExcel.ini')

    if not os.path.exists(config_file_path):
        print("INI файл не найден!")
        input("Нажмите Enter для выхода ...")
        exit(1)
    try:
        config.read(config_file_path, encoding='utf-8')
        template_name = config.get('PATHS', 'template_name')
        data_file_name = config.get('PATHS', 'data_file_name')

        # валидация
        if not template_name or not isinstance(template_name, str):
            raise ValueError("Название файла шаблона договора в ini файле указано не корректно")
        if not data_file_name or not isinstance(data_file_name, str):
            raise ValueError("Название файла эксел с данными в ini файле указано не корректно")
        if not template_name.lower().endswith('.docx'):
            raise ValueError("Название файла с шаблоном договора должен оканчиваться на .docx")
        if not data_file_name.lower().endswith('.xlsx'):
            raise ValueError("Название файла с данными экселе должен оканчиваться на .xlsx")
        if not os.path.exists(os.path.join(os.getcwd(), template_name)):
            raise ValueError(f"Не найден файл шаблона договора - {template_name}")
        if not os.path.exists(os.path.join(os.getcwd(), data_file_name)):
            raise ValueError(f"Не найден файл эксел с данными - {data_file_name}")

        return template_name, data_file_name

    except Exception as e:
        print(f"Ошибка: {e}.")
        print("Проверьте ini файл на корректное заполнение или убедитесь в наличии файла шаблона и файла с данными!")
        input("Нажмите Enter для выхода ...")
        exit(1)


def main():
    # Загрузка конфигурации
    template_name, data_file_name = load_config()
    # Конфигурация путей
    exe_dir = os.getcwd()

    template_path = os.path.join(exe_dir, template_name)
    xlsx_path = os.path.join(exe_dir, data_file_name)

    try:
        # Работа с Excel-данными
        wb = openpyxl.load_workbook(xlsx_path)
        sheet = wb.active
        rows = list(sheet.iter_rows(values_only=True))

        if not rows:
            raise ValueError("Файл Excel не содержит данных")

        headers = []
        for cell in rows[0]:
            if cell is not None and str(cell).strip() != "":
                headers.append(str(cell))
            else:
                break
        columns_count = len(headers)
        if columns_count == 0:
            raise ValueError("Все ячейки верхней строки Excel файла должны быть заполнены!")

        for row_idx, row in enumerate(rows[1:], 1):
            # Нормализация данных
            row_data = []
            for cell in row[:columns_count]:
                if cell is None:
                    row_data.append("-")
                elif isinstance(cell, datetime):
                    row_data.append(cell.strftime("%d.%m.%Y"))
                else:
                    row_data.append(str(cell).strip() or "-")

            # Дополнение данных до количества столбцов
            row_data += ["-"] * (columns_count - len(row_data))

            # Извлечение имени документа
            doc_name = row_data[0].strip() or f"row_{row_idx}"
            if not doc_name:
                print(f"Предупреждение: Пустое имя в строке {row_idx}, пропуск")
                continue

            # Загрузка шаблона
            doc = Document(template_path)

            # Выполнение замен
            for col_idx in range(1, columns_count):
                placeholder = headers[col_idx]
                value = row_data[col_idx]
                replace_text_in_doc(doc, placeholder, value)
                print(f"Замена: {placeholder} → {value}")

            # Проверка незамененных плейсхолдеров
            doc_text = get_document_text(doc)
            missing = [ph for ph in headers[1:] if ph in doc_text]

            if missing:
                missing_str = ", ".join(missing)
                print(f"Незамененные значения в шаблоне: {missing_str}.")
                error_msg = (
                    "Незамененные значения в шаблоне: "
                    f"{missing_str}. Убедитесь в едином форматировании!"
                )

                ps_script = (
                    'Add-Type -AssemblyName PresentationFramework;'
                    f'[System.Windows.MessageBox]::Show("{error_msg}", "Ошибка")'
                )
                subprocess.run(
                    ["powershell", "-Command", ps_script],
                    check=False
                )

            # Сохранение результата
            output_file = f"{doc_name}{Path(template_name).suffix}"
            output_path = os.path.join(exe_dir, output_file)
            doc.save(output_path)
            print(f"Создан документ: {output_path}\n")

        wb.close()

    except Exception as e:
        print(f"КРИТИЧЕСКАЯ ОШИБКА: {str(e)}", file=sys.stderr)
        input("Нажмите Enter для выхода ...")
        sys.exit(1)

if __name__ == "__main__":
    main()